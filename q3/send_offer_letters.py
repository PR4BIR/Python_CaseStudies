import pandas as pd
import smtplib
from email.message import EmailMessage
from email.utils import make_msgid
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
import os
import mimetypes
import logging
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASS = os.getenv("EMAIL_PASS")
COMPANY_NAME = os.getenv("COMPANY_NAME")

# Logo image path
logo_path = r"C:/HTC-2025/PYTHON/EXCERCISE/q3/logo.png"

# Set up logging
logging.basicConfig(filename='email_errors.log', level=logging.ERROR,
                    format='%(asctime)s:%(levelname)s:%(message)s')

# Read Excel file
try:
    df = pd.read_excel("C:/HTC-2025/PYTHON/EXCERCISE/q3/job_offers.xlsx")
except Exception as e:
    logging.error(f"Error reading Excel file: {e}")
    raise SystemExit("Failed to read Excel file.")

# Function to send email
def send_email(candidate):
    try:
        name = candidate["CandidateName"]
        email = candidate["Email"]
        role = candidate["JobRole"]
        joining = candidate["JoiningDate"]
        ctc = candidate["CTC"]
        candidate_id = candidate["CandidateID"]

        # Create offer letter docx and pdf filenames
        docx_filename = f"offer_{candidate_id}.docx"
        pdf_filename = f"offer_{candidate_id}.pdf"

        doc = Document()
        
        # Add logo at the top of the offer letter
        doc.add_picture(logo_path, width=Inches(2))  # Adjust size as needed
        doc.add_paragraph("\n")  # Space after logo
        
        doc.add_paragraph(f"Dear {name},\n")
        doc.add_paragraph(f"We are pleased to offer you the position of {role} at {COMPANY_NAME}.\n")
        doc.add_paragraph(f"Your joining date is {joining} and your CTC will be ₹{ctc}.\n")
        doc.add_paragraph("Kindly confirm your acceptance by replying to this email.\n")
        doc.add_paragraph("We look forward to welcoming you on board.\n")
        doc.add_paragraph("\nBest Regards,\nHR Team\n" + COMPANY_NAME)
        doc.save(docx_filename)

        # Convert DOCX to PDF
        convert(docx_filename, pdf_filename)

        # Create email message
        msg = EmailMessage()
        msg['Subject'] = f"Congratulations {name}! Your Job Offer from {COMPANY_NAME}"
        msg['From'] = EMAIL_USER
        msg['To'] = email

        # Embed logo in email HTML
        logo_cid = make_msgid(domain="htcglobal.com")[1:-1]

        # Email body (HTML)
        html = f"""
        <html>
            <body>
                <img src="cid:{logo_cid}" width="200"><br><br>
                <p>Dear {name},</p>
                <p>We are pleased to offer you the position of <b>{role}</b> at <b>{COMPANY_NAME}</b>.</p>
                <p>Your joining date is <b>{joining}</b> and your CTC will be ₹<b>{ctc}</b>.</p>
                <p>Kindly confirm your acceptance by replying to this email.</p>
                <p>We look forward to welcoming you on board.</p>
                <br>
                <p>Best Regards,<br>HR Team<br>{COMPANY_NAME}</p>
            </body>
        </html>
        """
        msg.set_content("Please view this email in HTML format.")
        msg.add_alternative(html, subtype='html')

        # Attach logo image inline in the email
        with open(logo_path, 'rb') as img:
            img_data = img.read()
            maintype, subtype = mimetypes.guess_type(logo_path)[0].split('/')
            msg.get_payload()[1].add_related(img_data, maintype=maintype, subtype=subtype, cid=logo_cid)

        # Attach PDF offer letter
        with open(pdf_filename, 'rb') as f:
            msg.add_attachment(f.read(), maintype='application', subtype='pdf', filename=pdf_filename)

        # Send the email via SMTP_SSL
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(EMAIL_USER, EMAIL_PASS)
            smtp.send_message(msg)

        print(f"✅ Email sent to {name} at {email}")

        # Cleanup files
        os.remove(docx_filename)
        os.remove(pdf_filename)

    except Exception as e:
        logging.error(f"Error sending email to {candidate['Email']}: {e}")
        print(f"❌ Error for {candidate['Email']}. See log.")

# Loop through candidates and send emails
for _, row in df.iterrows():
    send_email(row)
